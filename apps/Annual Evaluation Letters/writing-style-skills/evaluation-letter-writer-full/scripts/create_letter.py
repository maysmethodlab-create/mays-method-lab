"""
Create a formatted faculty evaluation letter as a .docx file.

This script produces a Word document with:
- Texas A&M University header logo (first page only)
- Proper memorandum format
- Consistent 11pt font throughout
- Page numbers on pages 2+
- Signature block at the end

Usage:
    python create_letter.py --json <path_to_letter_json> --output <output_path.docx> --logo <path_to_logo>

The JSON input file must have this structure:
{
    "date": "April 2026",
    "to_name": "Jon Jasperson",
    "to_title": "Associate Dean for Academic Innovation",
    "eval_year": "2025",
    "salutation_name": "Jon",
    "opening_paragraph": "Thank you for serving as the Associate Dean...",
    "accomplishments_paragraphs": [
        "Your contributions in 2025 reflect...",
        "You grew the Teaching & Learning Innovation team...",
        "In the classroom, you continued to model innovation..."
    ],
    "observations_paragraphs": [
        "Your performance this past year has been a model of...",
        "What stands out in our conversations is...",
        "As we continue to evolve, one area of growth..."
    ],
    "forward_looking_intro": "Your goals for 2026 are exactly what we need...",
    "forward_looking_bullets": [
        "Launch the Mays Flex XR Stage and shepherd its implementation...",
        "Embed XR-enabled redesigns into each Flex Online block...",
        "Recruit strategically to support the expanding infrastructure..."
    ],
    "forward_looking_closing": "This thoughtful and ambitious roadmap will ensure...",
    "summary_paragraphs": [
        "Jon, this was a landmark year. You made bold moves...",
        "Overall, my evaluation is that you have demonstrated excellent performance. Please return a signed copy of this annual performance review for our personnel files. Thank you."
    ]
}
"""

import json
import argparse
import os
import sys

def create_letter(letter_data, output_path, logo_path):
    """Create the evaluation letter .docx file."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
        from docx import Document
        from docx.shared import Inches, Pt, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

    doc = Document()

    # =========================================
    # PAGE SETUP
    # =========================================
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    # Enable "Different first page" for header/footer
    section.different_first_page_header_footer = True

    # =========================================
    # FIRST PAGE HEADER (with A&M logo)
    # =========================================
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False

    if logo_path and os.path.exists(logo_path):
        header_para = first_header.paragraphs[0] if first_header.paragraphs else first_header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(4.5))
        # Add a bit of space after the logo
        header_para_format = header_para.paragraph_format
        header_para_format.space_after = Pt(6)

    # Default header (pages 2+) - empty
    default_header = section.header
    default_header.is_linked_to_previous = False
    if default_header.paragraphs:
        default_header.paragraphs[0].text = ""

    # =========================================
    # FOOTER (page numbers on pages 2+ only)
    # =========================================
    # First page footer - empty
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if first_footer.paragraphs:
        first_footer.paragraphs[0].text = ""

    # Default footer (pages 2+) - page number
    default_footer = section.footer
    default_footer.is_linked_to_previous = False
    footer_para = default_footer.paragraphs[0] if default_footer.paragraphs else default_footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = footer_para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = footer_para.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

    for run in footer_para.runs:
        run.font.size = Pt(11)

    # =========================================
    # HELPER FUNCTIONS
    # =========================================
    def set_font(run, size=11, bold=False, font_name=None):
        """Set font properties on a run."""
        run.font.size = Pt(size)
        run.bold = bold
        if font_name:
            run.font.name = font_name

    def add_body_paragraph(text, space_after=Pt(12)):
        """Add a body text paragraph with consistent formatting."""
        para = doc.add_paragraph()
        para.paragraph_format.space_after = space_after
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(text)
        set_font(run, size=11)
        return para

    def add_bold_heading(text, space_before=Pt(18), space_after=Pt(6)):
        """Add a bold section heading (not a Word heading style, just bold body text)."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = space_before
        para.paragraph_format.space_after = space_after
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(text)
        set_font(run, size=11, bold=True)
        return para

    def add_memo_line(label, value, label_bold=True):
        """Add a TO/FROM/SUBJECT line with bold label."""
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.15

        run_label = para.add_run(label)
        set_font(run_label, size=11, bold=label_bold)

        run_value = para.add_run(value)
        set_font(run_value, size=11)
        return para

    def add_memo_continuation(text):
        """Add a continuation line (indented, for title under name)."""
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        # Add tab for indentation to align with content after label
        run = para.add_run("\t" + text)
        set_font(run, size=11)
        return para

    # =========================================
    # DOCUMENT CONTENT
    # =========================================

    # Date
    date_para = add_body_paragraph(letter_data["date"], space_after=Pt(18))

    # MEMORANDUM heading
    memo_heading = doc.add_paragraph()
    memo_heading.paragraph_format.space_after = Pt(18)
    memo_heading.paragraph_format.space_before = Pt(0)
    run = memo_heading.add_run("MEMORANDUM")
    set_font(run, size=11, bold=True)

    # TO line
    add_memo_line("TO:\t", letter_data["to_name"])
    if letter_data.get("to_title"):
        add_memo_continuation(letter_data["to_title"])

    # Blank line
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)

    # FROM line
    add_memo_line("FROM:\t", "Shrihari (Hari) Sridhar, Ph.D.")
    add_memo_continuation("Senior Associate Dean, Mays Business School")

    # Blank line
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)

    # SUBJECT line
    add_memo_line("SUBJECT:\t", f"{letter_data['eval_year']} Performance Evaluation")

    # Blank line before salutation
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    spacer.paragraph_format.space_before = Pt(6)

    # Salutation
    add_body_paragraph(f"Dear {letter_data['salutation_name']},")

    # Opening paragraph
    add_body_paragraph(letter_data["opening_paragraph"])

    # Summary of Major Accomplishments
    add_bold_heading("Summary of Major Accomplishments")
    for para_text in letter_data["accomplishments_paragraphs"]:
        add_body_paragraph(para_text)

    # My Observations and Our Discussion
    add_bold_heading("My Observations and Our Discussion")
    for para_text in letter_data["observations_paragraphs"]:
        add_body_paragraph(para_text)

    # Your Plan for the Upcoming Year
    add_bold_heading("Your Plan for the Upcoming Year")

    # Intro sentence
    if letter_data.get("forward_looking_intro"):
        add_body_paragraph(letter_data["forward_looking_intro"])

    # Bullet points for goals
    for bullet_text in letter_data.get("forward_looking_bullets", []):
        bullet_para = doc.add_paragraph(style='List Bullet')
        bullet_para.paragraph_format.space_after = Pt(4)
        bullet_para.paragraph_format.space_before = Pt(2)
        bullet_para.paragraph_format.line_spacing = 1.15
        # Clear existing runs and add formatted one
        bullet_para.clear()
        run = bullet_para.add_run(bullet_text)
        set_font(run, size=11)

    # Closing sentence after bullets
    if letter_data.get("forward_looking_closing"):
        add_body_paragraph(letter_data["forward_looking_closing"])

    # Summary
    add_bold_heading("Summary")
    for para_text in letter_data["summary_paragraphs"]:
        # Check if this is the evaluation statement (make it bold)
        if "Overall, my evaluation" in para_text:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            run = para.add_run(para_text)
            set_font(run, size=11, bold=True)
        else:
            add_body_paragraph(para_text)

    # Signature block
    sig_spacer = doc.add_paragraph()
    sig_spacer.paragraph_format.space_before = Pt(24)
    sig_spacer.paragraph_format.space_after = Pt(6)
    run = sig_spacer.add_run("_" * 55)
    set_font(run, size=11)

    sig_line = doc.add_paragraph()
    sig_line.paragraph_format.space_after = Pt(0)
    run = sig_line.add_run("Signature\t\t\t\t\t\tDate")
    set_font(run, size=11)

    # =========================================
    # SAVE
    # =========================================
    doc.save(output_path)
    print(f"Letter saved to: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Create a faculty evaluation letter .docx")
    parser.add_argument("--json", required=True, help="Path to the JSON file with letter content")
    parser.add_argument("--output", required=True, help="Output path for the .docx file")
    parser.add_argument("--logo", default=None, help="Path to the Texas A&M header logo image")
    args = parser.parse_args()

    with open(args.json, 'r') as f:
        letter_data = json.load(f)

    create_letter(letter_data, args.output, args.logo)


if __name__ == "__main__":
    main()
